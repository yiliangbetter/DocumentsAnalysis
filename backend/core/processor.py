"""Document processing - text extraction and chunking."""
import re
from io import BytesIO
from pathlib import Path
from typing import BinaryIO, List, NamedTuple, Optional, Tuple

from config import settings
from core.document import (
    Document,
    DocumentChunk,
    DocumentMetadata,
    DocumentType,
    ProcessingStatus,
)


class ProcessingResult(NamedTuple):
    """Result of processing a document."""
    document: Document
    chunks: List[DocumentChunk]


class DocumentProcessor:
    """Process documents - extract text, metadata, and create chunks."""

    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP

    def process_document(
        self,
        file_content: bytes,
        filename: str,
        doc_type: DocumentType,
    ) -> ProcessingResult:
        """Process a document file and return a ProcessingResult."""
        # Create document object
        document = Document(
            title=Path(filename).stem,
            source_path=filename,
            doc_type=doc_type,
            status=ProcessingStatus.PROCESSING,
        )

        try:
            # Extract text and metadata based on document type
            text, metadata = self._extract_content(file_content, doc_type)

            # Update document metadata
            document.metadata = metadata
            if metadata.title:
                document.title = metadata.title

            # Create chunks
            chunks = self._create_chunks(text, document.id)
            document.chunk_count = len(chunks)

            # Mark as completed
            document.status = ProcessingStatus.COMPLETED

            return ProcessingResult(document, chunks)

        except Exception as e:
            document.status = ProcessingStatus.FAILED
            document.error_message = str(e)
            return ProcessingResult(document, [])

    def _extract_content(
        self,
        file_content: bytes,
        doc_type: DocumentType,
    ) -> Tuple[str, DocumentMetadata]:
        """Extract text and metadata based on document type."""
        if doc_type == DocumentType.PDF:
            return self._extract_pdf(file_content)
        elif doc_type == DocumentType.DOCX:
            return self._extract_docx(file_content)
        elif doc_type in [DocumentType.TXT, DocumentType.MD]:
            return self._extract_text(file_content)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

    def _extract_pdf(self, file_content: bytes) -> Tuple[str, DocumentMetadata]:
        """Extract text from PDF using configured parser."""
        from pypdf import PdfReader

        metadata = DocumentMetadata()
        text_parts = []

        try:
            reader = PdfReader(BytesIO(file_content))

            # Extract metadata
            if reader.metadata:
                meta = reader.metadata
                metadata.title = meta.get("/Title", "")
                metadata.author = meta.get("/Author", "")
                metadata.subject = meta.get("/Subject", "")
                metadata.creator = meta.get("/Creator", "")
                metadata.creation_date = meta.get("/CreationDate", "")
                metadata.modification_date = meta.get("/ModDate", "")

            metadata.page_count = len(reader.pages)

            # Extract text from each page
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            full_text = "\n\n".join(text_parts)
            metadata.word_count = len(full_text.split())
            metadata.file_size = len(file_content)

            return full_text, metadata

        except Exception as e:
            raise ValueError(f"Failed to extract PDF content: {str(e)}")

    def _extract_docx(self, file_content: bytes) -> Tuple[str, DocumentMetadata]:
        """Extract text from DOCX."""
        from docx import Document as DocxDocument

        metadata = DocumentMetadata()

        try:
            doc = DocxDocument(BytesIO(file_content))

            # Extract metadata from core properties
            core_props = doc.core_properties
            metadata.title = core_props.title or ""
            metadata.author = core_props.author or ""
            metadata.subject = core_props.subject or ""
            metadata.creation_date = str(core_props.created) if core_props.created else None
            metadata.modification_date = str(core_props.modified) if core_props.modified else None

            # Extract text from paragraphs
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(text_parts)

            metadata.word_count = len(full_text.split())
            metadata.file_size = len(file_content)

            return full_text, metadata

        except Exception as e:
            raise ValueError(f"Failed to extract DOCX content: {str(e)}")

    def _extract_text(self, file_content: bytes) -> Tuple[str, DocumentMetadata]:
        """Extract text from plain text files."""
        metadata = DocumentMetadata()

        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = file_content.decode("utf-8")
            except UnicodeDecodeError:
                text = file_content.decode("latin-1")

            metadata.word_count = len(text.split())
            metadata.file_size = len(file_content)

            return text, metadata

        except Exception as e:
            raise ValueError(f"Failed to extract text content: {str(e)}")

    def _create_chunks(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        chunks = []
        text_length = len(text)

        # Simple character-based chunking with overlap
        start = 0
        chunk_index = 0

        while start < text_length:
            end = min(start + self.chunk_size, text_length)

            # Try to break at sentence or word boundary
            if end < text_length:
                # Look for sentence ending
                next_period = text.find(". ", end - 50, end + 50)
                if next_period != -1 and next_period < end + 50:
                    end = next_period + 1
                else:
                    # Break at word boundary
                    while end > start and not text[end - 1].isspace():
                        end -= 1
                    if end == start:
                        end = min(start + self.chunk_size, text_length)

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk = DocumentChunk(
                    document_id=document_id,
                    text=chunk_text,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end,
                )
                chunks.append(chunk)
                chunk_index += 1

            # Move start with overlap
            start = end - self.chunk_overlap
            if start >= end:
                start = end

        return chunks


def get_document_type(filename: str) -> DocumentType:
    """Determine document type from filename extension."""
    ext = Path(filename).suffix.lower()
    type_map = {
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".txt": DocumentType.TXT,
        ".md": DocumentType.MD,
        ".markdown": DocumentType.MD,
    }
    return type_map.get(ext, DocumentType.UNKNOWN)
